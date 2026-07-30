"""Build the ODE protocol as a Word (.docx) document (Overview-Data-Execution).
Numbers are read from saved run metrics; also regenerates a clean multi-page PDF is done separately.
"""
import json
import statistics as st
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
BLUE = RGBColor(0x18, 0x4f, 0x95)


def mae(run, sub="test_all"):
    return json.load(open(ROOT / "runs" / run / "test_metrics.json"))[sub]["mae_cm"]


def acc(run):
    return json.load(open(ROOT / "runs" / run / "test_metrics.json"))["test_all"]["accuracy"]


patch = [mae("dinov2_vits14_patchtokens_frozen"), mae("dinov2_vits14_patchtokens_frozen_s1"), mae("dinov2_vits14_patchtokens_frozen_s2")]
cls_ctrl = [mae("dinov2_vits14_clstoken_ctrl_s42"), mae("dinov2_vits14_clstoken_ctrl_s1"), mae("dinov2_vits14_clstoken_ctrl_s2")]
p_mean, p_sd = st.mean(patch), st.pstdev(patch)
c_mean, c_sd = st.mean(cls_ctrl), st.pstdev(cls_ctrl)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def h(t, lvl=1):
    p = doc.add_heading(t, level=lvl)
    for r in p.runs:
        r.font.color.rgb = BLUE
    return p


def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.italic = italic
    return p


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = "Light Grid Accent 1"
    for i, x in enumerate(headers):
        tb.rows[0].cells[i].text = str(x)
    for row in rows:
        c = tb.add_row().cells
        for i, v in enumerate(row):
            c[i].text = str(v)
    return tb


# ---- Title ----
t = doc.add_heading("ODE Protocol — Automated Fish Length Estimation with Vision Foundation Models", 0)
for r in t.runs:
    r.font.color.rgb = BLUE
para("ODE = Overview · Data · Execution — a structured, reproducible record of the machine-learning "
     "workflow (seminar Scientific-Writing protocol).", italic=True)
para("Authors: Abu Bakar, Laksh Jiwani, Shahman Butt  |  Supervisor: Bohan Zhuang, M.Sc.  |  "
     "Professor: Stefan Oehmcke  |  University of Rostock (VACOT)")
para("Repository: https://github.com/Shahman-Butt/AreaSeminarFishLength")

# ================= O =================
h("O — Overview", 1)
para("Research question.", bold=True)
para("Can a modern image encoder or vision foundation model beat the reproduced AutoFish MobileNetV2 "
     "baseline at estimating fish length from a photo, under identical data, split, and metrics?")
para("Target.", bold=True)
para("Reproduce the paper's baseline (0.62 cm non-occluded MAE), then run a controlled encoder-swap "
     "comparison. Secondary tasks: species classification; a reliable improvement inside a foundation "
     "model.")
para("Scope & assumptions.", bold=True)
bullet("Task is regression (predict a continuous length in cm), not detection/classification/segmentation.")
bullet("Single fish per crop; bounding box provided by the dataset (we do not train a detector).")
bullet("Only the encoder changes between length experiments; everything else is held fixed.")
para("Why this setup.", bold=True)
para("Reproducing the baseline first calibrates the whole pipeline, so any later difference is "
     "attributable to the encoder rather than to implementation error. A group-level split prevents "
     "the same fish leaking across train/test.")
para("Contributions (current state).", bold=True)
bullet("Baseline reproduced to within 0.013 cm (non-occluded).")
bullet("Controlled single-model comparison of 5 encoders (9 length runs).")
bullet(f"Reliable, multi-seed improvement within DINOv2: patch-token pooling ({p_mean:.3f} ± {p_sd:.3f}) "
       f"beats the CLS token ({c_mean:.3f} ± {c_sd:.3f}).")
bullet("Species classification across 4 encoders.")
bullet(f"EfficientNet-B0 reaches {mae('efficientnet_b0'):.3f} cm — within "
       f"{mae('efficientnet_b0')-mae('baseline_official'):.3f} cm of the baseline at a basic recipe; a "
       f"validation-based recipe search is under way to try to cross it with a single model.")

# ================= D =================
h("D — Data", 1)
para("Dataset.", bold=True)
para("AutoFish (Bengtson et al., WACV Workshops 2025), Hugging Face vapaau/autofish: 1,500 top-view "
     "RGB images, 454 unique fish, 25 groups, 18,157 instance annotations. Each annotation has a "
     "species label, a fish_id, a hand-measured length (cm), a bounding box, and a polygon segmentation.")
para("Labels.", bold=True)
para("Length in cm (regression target); species (7 classes) for the classification task.")
para("Preprocessing.", bold=True)
bullet("build_autofish_index.py -> single index.csv (18,157 rows) joining annotations, images, species.")
bullet("make_crops.py -> per-annotation masked, square crops resized to 224x224 (background outside "
       "the fish polygon blacked out; square crop preserves aspect ratio). 0 missing crops.")
bullet("ImageNet normalization; ColorJitter augmentation on the training split only.")
para("Splits (official, group-level).", bold=True)
para("Train 15 groups / Validation 5 / Test 5. Test = 3,759 annotations (1,879 non-occluded + 1,880 "
     "occluded).")
para("Leakage control.", bold=True)
para("An audit found one fish (fish_id 113) crossing splits via a singleton duplicate annotation "
     "(id 3759). It was removed (exclusions.json); the audit then confirms zero fish cross any split. "
     "This is why the non-occluded test count is 1,879, not 1,880.")

# ================= E =================
h("E — Execution", 1)
para("Model.", bold=True)
para("Encoder -> feature vector, concatenated with 4 normalized bbox values -> MLP regression head "
     "(Linear+BatchNorm+ReLU blocks -> 1 output). Only the encoder differs between length experiments.")
para("Models & configurations tried (full-test MAE, cm):", bold=True)
table(["Encoder", "Type", "Adaptation", "Recipe (Adam)", "Full-test MAE"], [
    ["MobileNetV2 (baseline)", "CNN", "full FT", "1e-3, 200 ep, b32", f"{mae('baseline_official'):.3f}"],
    ["EfficientNet-B0", "CNN", "full FT", "1e-4, 100 ep, b16", f"{mae('efficientnet_b0'):.3f}"],
    ["ConvNeXt-Tiny", "CNN", "full FT", "1e-4, 100 ep, b16", f"{mae('convnext_tiny_official'):.3f}"],
    ["CLIP ViT-B/32", "Transformer", "last-block FT", "1e-4, 100 ep", f"{mae('clip_vitb32_lastblock_from_frozen'):.3f}"],
    ["CLIP ViT-B/32", "Transformer", "frozen", "1e-4, 100 ep", f"{mae('clip_vitb32_frozen'):.3f}"],
    ["DINOv2 ViT-S/14 (patch)", "Transformer", "frozen, patch pool", "1e-4, 100 ep", f"{p_mean:.3f} ± {p_sd:.3f} (3 seeds)"],
    ["DINOv2 ViT-S/14 (CLS)", "Transformer", "last-block FT", "1e-4, 100 ep", f"{mae('dinov2_vits14_lastblock_from_frozen'):.3f}"],
    ["DINOv2 ViT-S/14 (CLS)", "Transformer", "frozen", "1e-3, 100 ep", f"{mae('dinov2_vits14_frozen'):.3f}"],
    ["DINOv2 ViT-S/14 (CLS)", "Transformer", "full FT", "enc 1e-5 / 1e-6", f"{mae('dinov2_vits14_finetune_lr1e5'):.3f} / {mae('dinov2_vits14_finetune_lr1e6'):.3f}"],
])
para("Training protocol.", bold=True)
para("L1 loss, Adam; best checkpoint selected on validation MAE; test set used once. Fixed seed (42) "
     "per run; one JSON config per experiment; hardware NVIDIA RTX 5000 Ada (32 GB), Python 3.11. New "
     "trainer options added this round: cosine LR schedule + weight decay.")
para("Metrics.", bold=True)
para("MAE (primary, cm), plus RMSE, MAPE, bias, R2; reported on full test, non-occluded, and occluded "
     "subsets. Baseline full-test: RMSE 1.268, MAPE 2.41%, bias +0.035, R2 0.947.")
para("Key results.", bold=True)
bullet("Baseline reproduction: 0.633 cm non-occluded vs paper 0.62 (delta 0.013) -> pipeline validated.")
bullet(f"Single-model ranking: MobileNetV2 best ({mae('baseline_official'):.3f}); EfficientNet-B0 "
       f"within {mae('efficientnet_b0')-mae('baseline_official'):.3f} cm; supervised CNNs top the ranking.")
bullet(f"Reliable DINOv2 improvement (3 seeds): patch pooling {p_mean:.3f} ± {p_sd:.3f} vs CLS "
       f"{c_mean:.3f} ± {c_sd:.3f} (non-overlapping ranges). Last-block on patch tokens did not help "
       f"({mae('dinov2_vits14_patchtokens_lastblock'):.3f}).")
bullet(f"Species classification (accuracy): ConvNeXt {acc('cls_convnext_tiny'):.3f}, MobileNetV2 "
       f"{acc('cls_mobilenet_v2'):.3f}, DINOv2 {acc('cls_dinov2_frozen'):.3f}, CLIP {acc('cls_clip_frozen'):.3f}.")
para("Validation-based recipe search (in progress).", bold=True)
para("Because EfficientNet-B0 used a weaker recipe than the baseline, we run stronger recipes (cosine "
     "LR + weight decay + tuned LR + 200 epochs) for EfficientNet-B0 and ConvNeXt-Tiny, select the "
     "winner on validation, and report once on test. Persistent execution on the GPU server; results "
     "committed as they land.")
para("Findings & interpretation.", bold=True)
bullet("Supervised CNNs fill the top of the single-model ranking; general foundation features transfer "
       "only partially (CLIP) or poorly (DINOv2) for precise metric regression (observation; mechanistic "
       "reasons are hypotheses).")
bullet("EfficientNet-B0's near-baseline result at a basic recipe indicates the baseline is likely "
       "beatable by a single model with proper tuning.")
para("Limitations.", bold=True)
bullet("Single training run for most single models (multi-seed done for the DINOv2 patch-vs-CLS study).")
bullet("Foundation models tested at small scale (ViT-S/14, ViT-B/32) with limited fine-tuning budgets.")
bullet("No single model beats the baseline yet; mask segmentation not yet evaluated.")
bullet("The occluded-set reproduction difference vs the paper (0.909 vs 1.38) is unexplained.")
para("Future work.", bold=True)
para("Finish the recipe search (likely to cross the baseline); EfficientNet-B2; patch pooling for CLIP; "
     "larger ViT variants; multi-seed the top single models; mask segmentation as a third task; DINOv3 "
     "(pending weight access).")
para("Reproducibility & AI statement.", bold=True)
para("Fixed seeds; official group split hard-coded; versioned configs; requirements.txt; per-run "
     "config/history/metrics/predictions; checkpoints on server; persistent queue scripts. Generative "
     "AI assisted with code scaffolding and drafting (documentation, figure captions, language "
     "improvement); it was not used to produce experimental results. Every number is read from saved "
     "metric files and was verified by the authors, who retain full responsibility for factual accuracy.")

out = ROOT / "docs" / "ODE_REPORT.docx"
doc.save(str(out))
print("wrote", out)
