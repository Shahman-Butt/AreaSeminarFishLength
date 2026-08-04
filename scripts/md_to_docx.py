"""Generic Markdown -> Word (.docx) converter for this project's docs.
Handles: #/##/### headers, **bold**, bullet lists, pipe tables, plain paragraphs, --- rules.
Usage: python scripts/md_to_docx.py <input.md> <output.docx> "<Title override or blank>"
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

BLUE = RGBColor(0x18, 0x4f, 0x95)


def add_richtext(paragraph, text):
    """Split on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            paragraph.add_run(part)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    if len(rows) >= 2:
        header = [c.strip() for c in rows[0].strip("|").split("|")]
        body = []
        for r in rows[2:]:
            cells = [c.strip() for c in r.strip("|").split("|")]
            body.append(cells)
        return header, body, i
    return None, None, i


def convert(md_path, docx_path, title_override=""):
    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.3)

    i = 0
    first_h1_done = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            txt = stripped[2:].strip()
            level = 0 if not first_h1_done else 1
            h = doc.add_heading(txt, level=0 if level == 0 else 1)
            for r in h.runs:
                r.font.color.rgb = BLUE
            first_h1_done = True
            i += 1
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:].strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = BLUE
            i += 1
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:].strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = BLUE
            i += 1
            continue

        if stripped.startswith("|"):
            header, body, new_i = parse_table(lines, i)
            if header:
                tb = doc.add_table(rows=1, cols=len(header))
                tb.style = "Light Grid Accent 1"
                for c, val in enumerate(header):
                    tb.rows[0].cells[c].text = val
                for row in body:
                    cells = tb.add_row().cells
                    for c, val in enumerate(row):
                        if c < len(cells):
                            cells[c].text = val
                i = new_i
                continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_richtext(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_richtext(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("*"))
            r.italic = True
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_richtext(p, stripped)
        i += 1

    doc.save(docx_path)
    print("wrote", docx_path)


if __name__ == "__main__":
    src = sys.argv[1]
    dst = sys.argv[2]
    convert(src, dst)
